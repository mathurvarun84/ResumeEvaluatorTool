"""
Verify Agent 4 merges gap-analysis sub_changes with ALL sectioner SubEntries.

Regression: fuzzy processed_labels dropped unchanged experience rows from DOCX output.
"""

from __future__ import annotations

import tempfile
from pathlib import Path
from unittest.mock import patch

from agents.rewriter import COMPANY_HEADER_START, RewriterAgent
from engine.resume_builder import build_final_docx
from schemas.common import SectionText, SubEntry
from validator.rewriter_validator import RewriterValidator


def _mock_llm_json(*_args, **_kwargs) -> str:
    """Minimal valid SectionRewrite JSON for one entry."""
    return (
        '{"balanced":"MockCo\\nRole\\n- Did things","aggressive":"MockCo\\nRole\\n- Did more",'
        '"top_1_percent":"MockCo\\nRole\\n- Peak"}'
    )


def test_ordered_merge_preserves_all_entries_when_labels_overlap() -> None:
    """Two roles at same company: only one targeted — both rows must appear."""
    entries = [
        SubEntry(
            label="Flipkart — EM (2021)",
            verbatim_text="Flipkart India\nEM\n- Led platform",
        ),
        SubEntry(
            label="Flipkart — Staff (2019)",
            verbatim_text="Flipkart India\nStaff\n- Earlier work",
        ),
        SubEntry(
            label="OtherCo — Dev (2018)",
            verbatim_text="OtherCo\nDev\n- Ship features",
        ),
    ]
    section_text = SectionText(
        header="experience",
        full_text="\n\n".join(e.verbatim_text for e in entries),
        sub_entries=list(entries),
    )
    gaps = [
        {
            "section": "experience",
            "needs_change": True,
            "rewrite_instruction": "Add metrics.",
            "sub_changes": [
                {
                    "sub_id": "1",
                    "sub_label": "Flipkart — EM (2021)",
                    "needs_change": True,
                    "rewrite_instruction": "Quantify impact.",
                    "missing_keywords": ["Python"],
                },
            ],
        }
    ]
    payload = {
        "gap_analysis": {"gaps": gaps, "strengths": [], "quick_wins": []},
        "resume_sections": {"experience": section_text.model_dump()},
        "resume_text": "",
        "jd_intelligence": None,
        "style_fingerprint": None,
    }

    agent = RewriterAgent()
    with patch.object(RewriterAgent, "_call_llm", side_effect=_mock_llm_json):
        out = agent.run(payload)

    balanced = out["rewrites"]["experience"]["balanced"]
    n_blocks = balanced.count(COMPANY_HEADER_START)
    assert n_blocks == len(entries), (
        f"expected {len(entries)} experience blocks, got {n_blocks}"
    )
    assert "Earlier work" in balanced, "verbatim Flipkart Staff row missing"
    assert "Ship features" in balanced, "verbatim OtherCo row missing"


def test_docx_contains_all_companies() -> None:
    """End-to-end: rewrites blob → build_final_docx lists every company."""
    entries = [
        SubEntry(
            label=f"Corp{i} — Role ({2015 + i})",
            verbatim_text=f"Corp{i}\nRole\n- Bullet {i}",
        )
        for i in range(8)
    ]
    section_text = SectionText(
        header="experience",
        full_text="\n\n".join(e.verbatim_text for e in entries),
        sub_entries=list(entries),
    )
    gaps = [
        {
            "section": "experience",
            "needs_change": True,
            "rewrite_instruction": "Improve.",
            "sub_changes": [
                {
                    "sub_id": "t0",
                    "sub_label": entries[0].label,
                    "needs_change": True,
                    "rewrite_instruction": "x",
                    "missing_keywords": [],
                },
                {
                    "sub_id": "t1",
                    "sub_label": entries[1].label,
                    "needs_change": True,
                    "rewrite_instruction": "y",
                    "missing_keywords": [],
                },
            ],
        }
    ]
    payload = {
        "gap_analysis": {"gaps": gaps, "strengths": [], "quick_wins": []},
        "resume_sections": {"experience": section_text.model_dump()},
        "resume_text": "",
        "jd_intelligence": None,
        "style_fingerprint": None,
    }

    agent = RewriterAgent()
    with patch.object(RewriterAgent, "_call_llm", side_effect=_mock_llm_json):
        out = agent.run(payload)

    structured = {"experience": section_text.model_dump()}
    doc = build_final_docx(structured, out["rewrites"], style="balanced")

    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
        path = Path(tmp.name)
    try:
        path.write_bytes(doc)
        from docx import Document

        full = "\n".join(p.text for p in Document(str(path)).paragraphs)
        for i in range(8):
            assert f"Corp{i}" in full, f"missing Corp{i} in DOCX body"
    finally:
        path.unlink(missing_ok=True)


def test_validator_repairs_partial_experience_rewrite_to_full_count() -> None:
    """Validator must append the 4 unchanged entries when A4 returns only 2."""
    entries = [
        SubEntry(
            label=f"Company{i} — Role ({2016 + i})",
            verbatim_text=f"Company{i}\nRole\n- Original bullet {i}",
        )
        for i in range(6)
    ]
    section_text = SectionText(
        header="experience",
        full_text="\n\n".join(e.verbatim_text for e in entries),
        sub_entries=entries,
    )
    partial = "\n\n".join(
        [
            f"{COMPANY_HEADER_START}Company0##ROLE##Role | 2016##END_HEADER##\n- Rewritten 0",
            f"{COMPANY_HEADER_START}Company1##ROLE##Role | 2017##END_HEADER##\n- Rewritten 1",
        ]
    )
    output = {
        "rewrites": {
            "experience": {
                "balanced": partial,
                "aggressive": partial,
                "top_1_percent": partial,
            }
        },
        "styles": {},
    }

    repaired = RewriterValidator().validate_and_fix(
        output,
        {"experience": section_text.model_dump()},
        "",
    )
    balanced = repaired["rewrites"]["experience"]["balanced"]

    assert balanced.count(COMPANY_HEADER_START) == 6
    assert "- Rewritten 0" in balanced
    assert "- Rewritten 1" in balanced
    for i in range(2, 6):
        assert f"Original bullet {i}" in balanced


def test_validator_does_not_duplicate_education_already_present() -> None:
    """Single-line education entries already present must not be appended again."""
    entries = [
        SubEntry(
            label="B.Tech Computer Science",
            verbatim_text="B.Tech Computer Science | VTU | 2015",
        ),
        SubEntry(
            label="MBA",
            verbatim_text="MBA | IIM Bangalore | 2020",
        ),
    ]
    full_text = "\n".join(e.verbatim_text for e in entries)
    section_text = SectionText(
        header="education",
        full_text=full_text,
        sub_entries=entries,
    )
    output = {
        "rewrites": {
            "education": {
                "balanced": full_text,
                "aggressive": full_text,
                "top_1_percent": full_text,
            }
        },
        "styles": {},
    }

    repaired = RewriterValidator().validate_and_fix(
        output,
        {"education": section_text.model_dump()},
        "",
    )
    balanced = repaired["rewrites"]["education"]["balanced"]

    assert balanced.count("B.Tech Computer Science") == 1
    assert balanced.count("MBA | IIM Bangalore") == 1


def test_validator_removes_duplicate_award_lines() -> None:
    """Flat award sections should not render duplicate entries."""
    awards = "Winner - Hackathon 2022\nWinner - Hackathon 2022\nTop 1% Performer"
    output = {
        "rewrites": {
            "awards": {
                "balanced": awards,
                "aggressive": awards,
                "top_1_percent": awards,
            }
        },
        "styles": {},
    }
    section_text = SectionText(
        header="awards",
        full_text="Winner - Hackathon 2022\nTop 1% Performer",
        sub_entries=[],
    )

    repaired = RewriterValidator().validate_and_fix(
        output,
        {"awards": section_text.model_dump()},
        "",
    )
    balanced = repaired["rewrites"]["awards"]["balanced"]

    assert balanced.count("Winner - Hackathon 2022") == 1
    assert "Top 1% Performer" in balanced


def test_validator_backfills_experience_from_resume_text() -> None:
    """If A1 has only 3 entries, validator should backfill from raw resume text."""
    entries = [
        SubEntry(
            label=f"Company{i} — Role ({2017 + i})",
            verbatim_text=f"Company{i}\nRole\n- Original bullet {i}",
        )
        for i in range(3)
    ]
    section_text = SectionText(
        header="experience",
        full_text="\n\n".join(e.verbatim_text for e in entries),
        sub_entries=entries,
    )
    partial = "\n\n".join(
        [
            f"{COMPANY_HEADER_START}Company0##ROLE##Role | 2017##END_HEADER##\n- Rewritten 0",
            f"{COMPANY_HEADER_START}Company1##ROLE##Role | 2018##END_HEADER##\n- Rewritten 1",
        ]
    )
    output = {
        "rewrites": {
            "experience": {
                "balanced": partial,
                "aggressive": partial,
                "top_1_percent": partial,
            }
        },
        "styles": {},
    }
    resume_text = """
EXPERIENCE
Company0 | Engineering Manager | 2017
- Delivered X

Company1 | Staff Engineer | 2018
- Delivered Y

Company2 | Senior Engineer | 2019
- Delivered Z

Company3 | Engineer | 2020
- Delivered A

Company4 | Engineer | 2021
- Delivered B

Company5 | Engineer | 2022
- Delivered C
"""
    repaired = RewriterValidator().validate_and_fix(
        output,
        {"experience": section_text.model_dump()},
        resume_text,
    )
    balanced = repaired["rewrites"]["experience"]["balanced"]
    assert balanced.count(COMPANY_HEADER_START) >= 6


def test_docx_skills_and_awards_render_as_list_items() -> None:
    """Skills comma-list and awards should render as list bullets."""
    structured = {
        "name": "Test User",
        "contact": "test@example.com",
        "resume_sections": {
            "skills": {"full_text": "", "sub_entries": []},
            "awards": {"full_text": "", "sub_entries": []},
        },
    }
    rewrites = {
        "skills": {
            "balanced": "Python, FastAPI, React, PostgreSQL",
            "aggressive": "Python, FastAPI, React, PostgreSQL",
            "top_1_percent": "Python, FastAPI, React, PostgreSQL",
        },
        "awards": {
            "balanced": "Winner - Hackathon 2022; Top 1% Performer",
            "aggressive": "Winner - Hackathon 2022; Top 1% Performer",
            "top_1_percent": "Winner - Hackathon 2022; Top 1% Performer",
        },
    }

    docx_bytes = build_final_docx(structured=structured, rewrites=rewrites, style="balanced")
    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
        path = Path(tmp.name)
    try:
        path.write_bytes(docx_bytes)
        from docx import Document

        doc = Document(str(path))
        bullet_count = sum(
            1 for p in doc.paragraphs if p.style and p.style.name == "List Bullet"
        )
        assert bullet_count >= 4, f"expected bullets for skills/awards, got {bullet_count}"
    finally:
        path.unlink(missing_ok=True)


if __name__ == "__main__":
    test_ordered_merge_preserves_all_entries_when_labels_overlap()
    test_docx_contains_all_companies()
    test_validator_repairs_partial_experience_rewrite_to_full_count()
    test_validator_does_not_duplicate_education_already_present()
    test_validator_removes_duplicate_award_lines()
    test_validator_backfills_experience_from_resume_text()
    test_docx_skills_and_awards_render_as_list_items()
    print("OK — experience preservation tests passed")
