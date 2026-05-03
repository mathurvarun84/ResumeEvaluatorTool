"""
Deterministic resume builder for Gap Closer output.

Exact style match to reference resume:
- Name: Bold, 26pt, color 1F3864, centered
- Title: Italic, 12pt, color 555555, centered
- Contact: 9.5pt, color 555555, centered, links in 1F5FA6
- Section headers: Bold, 11pt, CAPS, color 1F3864, w:spacing=40
- Ruled line below header: bottom border, w:sz=8, color 1F3864
- Body text: Arial 9.5pt, color 2E2E2E
- Role titles: Bold 1F3864 11pt
"""

from __future__ import annotations

import re
from io import BytesIO

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt

# ── Color palette (from reference resume docx XML) ──
COLOR_NAME       = "1F3864"   # dark navy — candidate name + role/company
COLOR_TITLE      = "555555"   # grey — subtitle/contact/meta
COLOR_SECTION_HDR= "1F3864"   # dark navy — section header text + border
COLOR_BODY       = "2E2E2E"   # near-black — body text and bullets
COLOR_LINK       = "1F5FA6"   # blue — hyperlinks in contact
COLOR_DATE       = "888888"   # light grey italic — date range
COLOR_ROLE_CO    = COLOR_NAME  # role/company blue

# ── Section canonical order and labels ──
SECTION_ORDER = ["summary", "skills", "experience", "education", "certifications", "awards", "projects"]
SECTION_LABELS = {
    "summary": "SUMMARY", "skills": "SKILLS", "experience": "EXPERIENCE",
    "education": "EDUCATION", "certifications": "CERTIFICATIONS", "awards": "AWARDS",
    "projects": "PROJECTS",
}

# ── Section name normalization ──
SECTION_NAME_MAP = {
    "professional_summary": "summary",
    "objective": "summary",
    "profile": "summary",
    "summary": "summary",
    "core_competencies": "skills",
    "key_skills": "skills",
    "skills": "skills",
    "technical_skills": "skills",
    "work_experience": "experience",
    "professional_experience": "experience",
    "employment": "experience",
    "experience": "experience",
    "education": "education",
    "certifications": "certifications",
    "certificates": "certifications",
    "awards": "awards",
    "achievements": "awards",
    "projects": "projects",
    "languages": "languages",
    "interests": "interests",
}


def _normalize_key(key: str) -> str:
    """Map variant section keys to canonical names."""
    return SECTION_NAME_MAP.get(key, key)


def _set_color(run, hex_color: str):
    """Set font color on a docx run via w:color element."""
    r_pr = run._r.get_or_add_rPr()
    color_el = OxmlElement("w:color")
    color_el.set(qn("w:val"), hex_color)
    r_pr.append(color_el)


def build_final_docx(structured_resume: dict, rewrites: dict, style: str = "balanced") -> bytes:
    """
    Build a .docx resume with exact style matching.

    Args:
        structured_resume: Dict from parser._build_structured_resume with
            'name', 'contact', and section keys (summary/skills/experience/etc.).
        rewrites: Dict {section_key: {style_str: content}} where section keys
            may be canonical or aliased (e.g. 'work_experience' → 'experience').
        style: Which rewrite style to use ('balanced', 'aggressive', 'top_1_percent').

    Returns:
        BytesIO.getvalue() of the .docx file.
    """
    # Normalize rewrite keys to canonical names
    normalized_rewrites = {}
    for key, val in rewrites.items():
        canonical = _normalize_key(key)
        normalized_rewrites[canonical] = val

    # Extract content for each section: rewrite takes priority, then structured
    sections = {}
    for key in SECTION_ORDER:
        rewrite = normalized_rewrites.get(key, {})
        if isinstance(rewrite, dict):
            content = rewrite.get(style, "")
        else:
            content = str(rewrite) if rewrite else ""

        # If rewrite is empty/unavailable, fall back to original structured content
        if not content or _is_unavailable_content(content):
            sections[key] = _extract_section_content(structured_resume, key)
        else:
            sections[key] = content

    candidate_name = structured_resume.get("name", "Candidate")
    contact_line = structured_resume.get("contact", "")

    return _build_docx(candidate_name, contact_line, sections, style)


def build_final_docx(structured: dict, rewrites: dict, style: str = "balanced") -> bytes:
    """Build a complete download-ready resume docx."""
    from docx import Document
    from docx.shared import Pt, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    from rewriter import COMPANY_HEADER_START, COMPANY_ROLE_START, HEADER_END
    import io, re

    doc = Document()
    for sec in doc.sections:
        sec.top_margin = sec.bottom_margin = Cm(1.6)
        sec.left_margin = sec.right_margin = Cm(1.9)

    name = (structured.get("name") or "").strip()
    title = (structured.get("title") or "").strip()
    contact = (structured.get("contact") or "").strip()

    if name:
        p = doc.add_paragraph()
        r = p.add_run(name.upper())
        r.bold = True
        r.font.size = Pt(18)
        r.font.color.rgb = RGBColor(0x1F, 0x38, 0x64)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if title:
        p = doc.add_paragraph()
        r = p.add_run(title)
        r.italic = True
        r.font.size = Pt(11)
        r.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if contact:
        p = doc.add_paragraph()
        r = p.add_run(contact)
        r.font.size = Pt(9)
        r.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    placeholder_re = re.compile(r'^\[.+\]$')
    year_re = re.compile(r'\b(19|20)\d{2}\b|[Pp]resent')
    bullet_markers = ('•', '-', '*')

    def _is_placeholder(t):
        return bool(placeholder_re.match(t.strip()))

    def _hdr(text):
        """Section header with bottom border, brand blue."""
        p = doc.add_paragraph()
        r = p.add_run(text.upper())
        r.bold = True
        r.font.size = Pt(11)
        r.font.color.rgb = RGBColor(0x1F, 0x38, 0x64)
        pPr = p._element.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bot = OxmlElement("w:bottom")
        bot.set(qn("w:val"), "single")
        bot.set(qn("w:sz"), "8")
        bot.set(qn("w:space"), "1")
        bot.set(qn("w:color"), "1F3864")
        pBdr.append(bot)
        pPr.append(pBdr)
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(2)

    def _company(company, location=""):
        """Bold company name line."""
        p = doc.add_paragraph()
        r = p.add_run(company)
        r.bold = True
        r.font.size = Pt(10)
        r.font.color.rgb = RGBColor(0x1F, 0x38, 0x64)
        if location:
            r2 = p.add_run(f"  {location}")
            r2.font.size = Pt(9)
            r2.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(1)

    def _role(role, dates=""):
        """Italic role + gray dates."""
        p = doc.add_paragraph()
        r = p.add_run(role)
        r.italic = True
        r.font.size = Pt(9.5)
        r.font.color.rgb = RGBColor(0x2E, 0x2E, 0x2E)
        if dates:
            r2 = p.add_run(f"  {dates}")
            r2.font.size = Pt(9)
            r2.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
        p.paragraph_format.space_after = Pt(1)

    def _bullet(text):
        """List Bullet paragraph with placeholder guard."""
        t = text.lstrip("-•* ").strip()
        if not t or _is_placeholder(t):
            return
        p = doc.add_paragraph(style="List Bullet")
        r = p.add_run(t)
        r.font.size = Pt(9.5)
        r.font.color.rgb = RGBColor(0x2E, 0x2E, 0x2E)
        p.paragraph_format.space_after = Pt(1)

    def _normal(text):
        """Normal paragraph with placeholder guard."""
        t = text.strip()
        if not t or _is_placeholder(t):
            return
        p = doc.add_paragraph()
        r = p.add_run(t)
        r.font.size = Pt(9.5)
        r.font.color.rgb = RGBColor(0x2E, 0x2E, 0x2E)
        p.paragraph_format.space_after = Pt(1)

    def _tech_stack(text):
        """Italic tech stack line."""
        p = doc.add_paragraph()
        r = p.add_run(text.strip())
        r.italic = True
        r.font.size = Pt(9)
        r.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
        p.paragraph_format.space_after = Pt(2)

    def _write_experience(content):
        """Write marker-based or plain experience with consistent formatting."""
        if COMPANY_HEADER_START in content:
            for block in content.split(COMPANY_HEADER_START):
                if not block.strip():
                    continue
                header_part = body = ""
                if HEADER_END in block:
                    header_part, body = block.split(HEADER_END, 1)
                else:
                    body = block
                co = loc = ro = dt = ""
                if COMPANY_ROLE_START in header_part:
                    co_loc, ro_dt = header_part.split(COMPANY_ROLE_START, 1)
                    co, loc = (co_loc.split("|", 1) + [""])[:2]
                    ro, dt = (ro_dt.split("|", 1) + [""])[:2]
                    co = co.strip()
                    loc = loc.strip()
                    ro = ro.strip()
                    dt = dt.strip()
                else:
                    co = header_part.strip()
                if co:
                    _company(co, loc)
                if ro:
                    _role(ro, dt)
                for line in body.splitlines():
                    s = line.strip()
                    if not s:
                        continue
                    if s.startswith(bullet_markers):
                        _bullet(s)
                    elif s.lower().startswith("tech stack:"):
                        _tech_stack(s)
                    else:
                        _normal(s)
        else:
            for raw in content.splitlines():
                s = raw.strip()
                if not s:
                    continue
                if s.startswith(bullet_markers):
                    _bullet(s)
                elif s.lower().startswith("tech stack:"):
                    _tech_stack(s)
                elif year_re.search(s):
                    _role(s)
                elif s[0].isupper() and len(s) < 80:
                    _company(s)
                else:
                    _normal(s)

    def _get_content(section_name):
        """Return rewrite content, balanced fallback, then structured original."""
        rw = rewrites.get(section_name, {})
        if not rw:
            # Accept aliased rewrite keys (e.g., work_experience -> experience)
            for key, value in rewrites.items():
                if _normalize_key(str(key)) == section_name:
                    rw = value
                    break
        for key in (style, "balanced"):
            val = rw.get(key, "") if isinstance(rw, dict) else ""
            if val and len(val.strip()) >= 10 and not _is_placeholder(val.strip()):
                return val
        return _extract_section_content(structured, section_name).strip()

    for sec_name in SECTION_ORDER:
        content = _get_content(sec_name)
        if not content:
            continue
        _hdr(sec_name)
        if sec_name == "experience":
            _write_experience(content)
        elif sec_name == "summary":
            p = doc.add_paragraph()
            r = p.add_run(content.strip())
            r.font.size = Pt(9.5)
            r.font.color.rgb = RGBColor(0x2E, 0x2E, 0x2E)
            p.paragraph_format.space_after = Pt(3)
        elif sec_name == "skills":
            for line in content.splitlines():
                s = line.strip()
                if not s or _is_placeholder(s):
                    continue
                if ":" in s:
                    p = doc.add_paragraph()
                    label, rest = s.split(":", 1)
                    rl = p.add_run(label + ":")
                    rl.bold = True
                    rl.font.size = Pt(9.5)
                    rl.font.color.rgb = RGBColor(0x1F, 0x38, 0x64)
                    rr = p.add_run(rest)
                    rr.font.size = Pt(9.5)
                    rr.font.color.rgb = RGBColor(0x2E, 0x2E, 0x2E)
                else:
                    # Convert comma/pipe-separated skill lists into bullets.
                    chunks = [
                        c.strip() for c in re.split(r"[|,;]\s*", s) if c.strip()
                    ]
                    if len(chunks) > 1:
                        for skill in chunks:
                            _bullet(f"• {skill}")
                    else:
                        p = doc.add_paragraph()
                        r = p.add_run(s)
                        r.font.size = Pt(9.5)
                        r.font.color.rgb = RGBColor(0x2E, 0x2E, 0x2E)
                        p.paragraph_format.space_after = Pt(1)
        elif sec_name == "awards":
            for line in content.splitlines():
                s = line.strip()
                if not s or _is_placeholder(s):
                    continue
                entries = [e.strip() for e in re.split(r";\s*", s) if e.strip()]
                if len(entries) > 1:
                    for entry in entries:
                        _bullet(f"• {entry}")
                elif s.startswith(bullet_markers):
                    _bullet(s)
                else:
                    _bullet(f"• {s}")
        else:
            for line in content.splitlines():
                s = line.strip()
                if not s or _is_placeholder(s):
                    continue
                if s.startswith(bullet_markers):
                    _bullet(s)
                else:
                    _normal(s)

    buf = io.BytesIO()
    doc.save(buf)
    result = buf.getvalue()
    assert len(result) > 5000, f"Docx too small: {len(result)} bytes"
    return result


def build_resume_docx(structured_resume: dict, rewrites: dict, style: str = "balanced") -> bytes:
    """Alias for build_final_docx (kept for backward compatibility)."""
    return build_final_docx(structured_resume, rewrites, style)


def _extract_section_content(structured: dict, section: str) -> str:
    """
    Extract text content for a section from structured resume dict.

    Handles lists (experience, education) and plain strings (summary, skills).
    Returns empty string if section not found.
    """
    if section == "experience":
        exp = _get_structured_value(structured, "experience")
        if isinstance(exp, list):
            items = []
            for item in exp:
                blocks = []
                for k in ("title", "company", "location", "dates"):
                    v = item.get(k, "")
                    if v:
                        blocks.append(str(v))
                for bullet in item.get("bullets", []) or []:
                    blocks.append(f"- {bullet}")
                if blocks:
                    items.append("\n".join(blocks))
            return "\n\n".join(items)
        return str(exp) if exp else ""

    if section == "education":
        edu = _get_structured_value(structured, "education")
        if isinstance(edu, list):
            lines = []
            for item in edu:
                row = " | ".join(
                    str(item.get(k, "")).strip()
                    for k in ("degree", "institution", "years")
                    if str(item.get(k, "")).strip()
                )
                if row:
                    lines.append(row)
            return "\n".join(lines)
        return str(edu) if edu else ""

    if section == "certifications":
        certs = _get_structured_value(structured, "certifications")
        if isinstance(certs, list):
            return "\n".join(str(c) for c in certs if c)
        return str(certs) if certs else ""

    if section == "awards":
        awards = _get_structured_value(structured, "awards")
        if isinstance(awards, list):
            return "\n".join(str(a) for a in awards if a)
        return str(awards) if awards else ""

    if section == "projects":
        projects = _get_structured_value(structured, "projects")
        if isinstance(projects, list):
            return "\n\n".join(str(project) for project in projects if project)
        return str(projects) if projects else ""

    # Simple string sections: summary, skills
    val = _get_structured_value(structured, section)
    return str(val) if val else ""


def _get_structured_value(structured: dict, section: str):
    """Read a canonical structured section, accepting common aliases."""
    if section in structured:
        return structured.get(section)
    for key, value in structured.items():
        if _normalize_key(key) == section:
            return value
    nested_sections = structured.get("resume_sections") or structured.get("sections") or {}
    if isinstance(nested_sections, dict):
        for key, value in nested_sections.items():
            if _normalize_key(str(key)) != section:
                continue
            if isinstance(value, dict):
                return value.get("full_text", "")
            if hasattr(value, "full_text"):
                return getattr(value, "full_text", "")
    return ""


def _is_unavailable_content(content) -> bool:
    """Detect internal placeholder strings that should not be rendered."""
    if not isinstance(content, str):
        return False
    lowered = content.strip().lower()
    return lowered.startswith("[") and (
        "unavailable" in lowered or "rewrite unavailable" in lowered
    )


def _build_docx(
    candidate_name: str,
    contact_line: str,
    sections: dict,
    style: str,
) -> bytes:
    """
    Build the .docx document with exact color/font styling.

    Args:
        candidate_name: Candidate's full name.
        contact_line: Contact info string (email | phone | LinkedIn | etc.).
        sections: Dict {canonical_section_name: text_content}.
        style: Rewrite style (unused in builder — content is already selected).

    Returns:
        BytesIO.getvalue() of the .docx file.
    """
    del style  # content is already selected for the correct style

    doc = Document()

    # ── Margins: top/bottom 900 twips (~1.58cm), left/right 1080 twips (~1.90cm) ──
    for sec in doc.sections:
        sec.top_margin = Pt(900 / 20)    # 900 twips / 20 = 45 pt
        sec.bottom_margin = Pt(900 / 20)
        sec.left_margin = Pt(1080 / 20)  # 1080 twips / 20 = 54 pt
        sec.right_margin = Pt(1080 / 20)

    # ── NAME: Bold, 26pt, color 1F3864, centered ──
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run((candidate_name or "Candidate").upper())
    r.font.name, r.font.bold = "Arial", True
    r.font.size = Pt(26)
    _set_color(r, COLOR_NAME)

    # ── TITLE: Italic, 12pt, color 555555, centered ──
    if structured_title := _detect_title(contact_line):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(3)
        r = p.add_run(structured_title)
        r.font.name, r.font.italic = "Arial", True
        r.font.size = Pt(12)
        _set_color(r, COLOR_TITLE)

    # ── CONTACT: 9.5pt, color 555555, centered ──
    if contact_line.strip():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(6)
        _add_contact_line(p, contact_line)

    # ── SECTIONS ──
    for key in SECTION_ORDER:
        content = sections.get(key, "")
        if not content or not content.strip():
            continue

        # Section header: Bold, 11pt, CAPS, color 1F3864, w:spacing=40
        hdr = doc.add_paragraph()
        hdr.paragraph_format.space_before = Pt(10)
        hdr.paragraph_format.space_after = Pt(2)
        label = SECTION_LABELS.get(key, key.upper())
        r = hdr.add_run(label)
        r.font.name, r.font.bold = "Arial", True
        r.font.size = Pt(11)
        _set_color(r, COLOR_SECTION_HDR)
        # Letter spacing: w:spacing val=40
        r_pr = r._r.get_or_add_rPr()
        spacing_el = OxmlElement("w:spacing")
        spacing_el.set(qn("w:val"), "40")
        r_pr.append(spacing_el)

        # Underline rule: bottom border on spacer paragraph
        rule = doc.add_paragraph()
        rule.paragraph_format.space_before = Pt(0)
        rule.paragraph_format.space_after = Pt(4)
        p_pr = rule._p.get_or_add_pPr()
        p_bdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "8")
        bottom.set(qn("w:space"), "2")
        bottom.set(qn("w:color"), COLOR_SECTION_HDR)
        p_bdr.append(bottom)
        p_pr.append(p_bdr)

        # Content
        if key == "summary":
            _render_summary(doc, content)
        elif key == "experience":
            if isinstance(content, str):
                _write_experience_section(doc, content)
            else:
                _render_experience(doc, content)
        else:
            _render_text_block(doc, content)

    buf = BytesIO()
    doc.save(buf)
    bval = buf.getvalue()
    assert len(bval) > 5000, f"Resume docx too small ({len(bval)} bytes) — likely missing sections"
    return bval


def _detect_title(contact_line: str) -> str:
    """Extract title from contact line if present."""
    return ""


def _add_contact_line(paragraph, text: str):
    """Add contact text with blue color for links (containing @ or URL-like tokens)."""
    for token in text.split(" " * 3):  # split on triple-space separators if present
        token_stripped = token.strip()
        if not token_stripped:
            continue
        is_link = "@" in token_stripped or "linkedin" in token_stripped.lower() or token_stripped.startswith("http")
        r = paragraph.add_run(token_stripped)
        r.font.name = "Arial"
        r.font.size = Pt(9.5)
        _set_color(r, COLOR_LINK if is_link else COLOR_TITLE)
        r = paragraph.add_run("   |   ")
        r.font.name = "Arial"
        r.font.size = Pt(9.5)
        _set_color(r, COLOR_TITLE)
    # Remove trailing " | " from last run
    if paragraph.runs and len(paragraph.runs) > 1:
        last_run = paragraph.runs[-1]
        if last_run.text == "   |   ":
            last_run.text = ""


def _render_summary(doc, content):
    """Render the summary as one paragraph instead of splitting by PDF lines."""
    text = str(content).strip()
    if not text:
        return
    p = doc.add_paragraph(text)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    for run in p.runs:
        run.font.name = "Arial"
        run.font.size = Pt(9.5)
        _set_color(run, COLOR_BODY)


def _write_experience_section(document, content: str):
    """
    Write flat rewritten experience text with company, role, bullet, and stack styles.
    """
    year_pattern = re.compile(r"\b(19|20)\d{2}\b|[Pp]resent")
    bullet_markers = ("-", "\u2022", "*")

    wrote_content = False
    pending_blank = False
    for raw_line in content.splitlines():
        line = raw_line.strip()
        if not line:
            if wrote_content:
                pending_blank = True
            continue

        if pending_blank:
            spacer = document.add_paragraph()
            spacer.paragraph_format.space_before = Pt(0)
            spacer.paragraph_format.space_after = Pt(2)
            pending_blank = False

        if line.startswith(bullet_markers):
            cleaned = line.lstrip("-\u2022* ").strip()
            p = document.add_paragraph(cleaned, style="List Bullet")
            _format_paragraph_runs(p)
        elif line.lower().startswith("tech stack:"):
            p = document.add_paragraph(line)
            _format_paragraph_runs(p, italic=True)
        elif year_pattern.search(line):
            p = document.add_paragraph(line)
            _format_paragraph_runs(p)
        else:
            p = document.add_paragraph(line)
            _format_paragraph_runs(p, bold=True)
        wrote_content = True


def _format_paragraph_runs(paragraph, bold: bool = False, italic: bool = False):
    paragraph.paragraph_format.space_before = Pt(2)
    paragraph.paragraph_format.space_after = Pt(2)
    for run in paragraph.runs:
        run.font.name = "Arial"
        run.font.size = Pt(9.5)
        run.font.bold = bold
        run.font.italic = italic
        _set_color(run, COLOR_BODY)


def _render_experience(doc, content):
    """
    Render experience section. Content can be a list of role dicts (verbatim)
    or a string (rewritten content).
    """
    if isinstance(content, str):
        _render_text_block(doc, content)
        return
    if isinstance(content, list):
        for role in content:
            # Role line: "Engineering Manager | Flipkart — Bengaluru, KA     Sep 2020 – Present"
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(2)

            r1 = p.add_run(f"{role.get('title', '')}  |  ")
            r1.font.name = "Arial"
            r1.font.size = Pt(10)
            r1.font.bold = True
            _set_color(r1, COLOR_ROLE_CO if 'COLOR_ROLE_CO' in globals() else COLOR_NAME)

            r2 = p.add_run(role.get("company", ""))
            r2.font.name = "Arial"
            r2.font.size = Pt(11)
            r2.font.bold = True
            _set_color(r2, COLOR_NAME)

            loc = role.get("location", "")
            if loc:
                r3 = p.add_run(f"  \u2014  {loc}")
                r3.font.name = "Arial"
                r3.font.size = Pt(10)
                _set_color(r3, COLOR_TITLE)

            dates = role.get("dates", "")
            if dates:
                r4 = p.add_run(f"\t{dates}")
                r4.font.name = "Arial"
                r4.font.size = Pt(9.5)
                r4.font.italic = True
                _set_color(r4, COLOR_DATE)

            for bullet in role.get("bullets", []) or []:
                bp = doc.add_paragraph(style="List Bullet")
                bp.paragraph_format.space_before = Pt(2)
                bp.paragraph_format.space_after = Pt(2)
                r = bp.add_run(bullet)
                r.font.name, r.font.size = "Arial", Pt(9.5)
                _set_color(r, COLOR_BODY)


def _render_text_block(doc, content):
    """
    Render a text block (summary, skills, education, etc.).
    Handles both string and list input.
    """
    if isinstance(content, list):
        items = content
    else:
        items = [line.strip() for line in str(content).split("\n") if line.strip()]

    for line in items:
        if not line:
            continue
        if line.startswith(("\u2022", "-", "*")):
            p = doc.add_paragraph(style="List Bullet")
            r = p.add_run(line.lstrip("\u2022-* "))
        else:
            p = doc.add_paragraph()
            r = p.add_run(line)
        r.font.name = "Arial"
        r.font.size = Pt(9.5)
        _set_color(r, COLOR_BODY)
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
