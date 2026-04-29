"""
Test script to verify ISSUE 1 and ISSUE 2 fixes for Gap Closer.

ISSUE 1: Verify DOCX export has all sections with fallback text
ISSUE 2: Verify Streamlit button handler has progress feedback
"""

import sys
import os
import tempfile
import logging
from pathlib import Path

# Setup logging to capture diagnostic output
logging.basicConfig(level=logging.DEBUG, format='%(name)s - %(levelname)s - %(message)s')
logger = logging.getLogger("test_gap_session_fixes")

print("=== ISSUE 1 & 2 FIX VERIFICATION ===\n")

failures = []
passes = []

def check(name, condition):
    """Helper to track pass/fail."""
    status = "PASS" if condition else "FAIL"
    print(f"  {status}: {name}")
    if condition:
        passes.append(name)
    else:
        failures.append(name)
    return condition

# ============================================================================
# ISSUE 1: DOCX Export with Null Guards and Fallback Text
# ============================================================================

print("1. ISSUE 1 - DOCX EXPORT WITH NULL GUARDS")
print("-" * 60)

try:
    from gap_session import run_gap_session, _export_to_docx, SectionResult
    from docx import Document

    check("gap_session module importable", True)
    check("_export_to_docx function exists", callable(_export_to_docx))
    check("SectionResult dataclass exists", True)

    # Create test data with some None/empty fields (simulating real scenario)
    test_full_text = "John Doe\njohn@example.com"
    test_sections = [
        SectionResult(
            name="Experience",
            original="Worked at Company",
            rewritten="Led team at Company",
            changes_made=["Added leadership keywords"],
            keywords_added=["Led", "team"],
            decision="accepted"
        ),
        SectionResult(
            name="Skills",
            original="",  # Empty original
            rewritten="Python, AWS",
            changes_made=["Added AWS"],
            keywords_added=["AWS"],
            decision="accepted"
        ),
    ]

    test_gap_result = {
        "target_role": "Senior Engineer",
        "jd_match_score_before": 45,
        "jd_match_score_after": 72,
        "gaps": [
            {"section": "Skills", "type": "Missing", "severity": "High", "suggestion": "Add cloud expertise"}
        ],
    }

    test_score_delta = {
        "keywords_added": ["Python", "AWS", "Leadership"],
        "sections_improved": ["Skills", "Experience"],
        "remaining_gaps": ["Need DevOps experience"],
        "manual_suggestions": ["Add Kubernetes project"],
    }

    # Generate DOCX to temp file
    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
        tmp_path = tmp.name

    try:
        _export_to_docx(
            full_text=test_full_text,
            sections=test_sections,
            path=tmp_path,
            gap_result=test_gap_result,
            score_delta=test_score_delta,
        )

        check("DOCX file created successfully", os.path.exists(tmp_path))

        file_size = os.path.getsize(tmp_path)
        logger.debug(f"Generated DOCX size: {file_size} bytes")
        check(f"DOCX file size > 5000 bytes (actual: {file_size})", file_size > 5000)

        # Verify content by opening the DOCX
        try:
            doc = Document(tmp_path)

            # Count paragraphs and tables
            para_count = len(doc.paragraphs)
            table_count = len(doc.tables)
            logger.debug(f"DOCX contains {para_count} paragraphs, {table_count} tables")

            # Collect all text to verify sections exist
            all_text = "\n".join([p.text for p in doc.paragraphs])

            # Verify key sections are present
            section_checks = [
                ("Header with name", "John Doe" in all_text),
                ("Target role in header", "Target Role" in all_text or "Senior Engineer" in all_text),
                ("UPGRADED RESUME section", "UPGRADED" in all_text),
                ("GAP ANALYSIS section", "GAP ANALYSIS" in all_text),
                ("JD MATCH ANALYSIS section", "JD MATCH" in all_text or "BEFORE" in all_text),
                ("IMPROVEMENT SUMMARY section", "IMPROVEMENT" in all_text or "Keywords Added" in all_text),
                ("Experience section heading", "EXPERIENCE" in all_text),
                ("Skills section heading", "SKILLS" in all_text),
            ]

            for check_name, result in section_checks:
                check(f"  {check_name}", result)

            # Verify no silent failures (empty content)
            check("GAP ANALYSIS table not empty", table_count > 0)
            check("Document has reasonable content (>30 paragraphs)", para_count > 30)

        except Exception as e:
            logger.error(f"Error reading DOCX content: {e}")
            check("Can open and read DOCX content", False)

    finally:
        # Cleanup
        if os.path.exists(tmp_path):
            os.unlink(tmp_path)

except ImportError as e:
    logger.error(f"Import error: {e}")
    check("gap_session module imports", False)
except Exception as e:
    logger.error(f"Unexpected error in ISSUE 1 test: {e}")
    import traceback
    traceback.print_exc()
    check("ISSUE 1 test completed without exception", False)

# ============================================================================
# ISSUE 2: Streamlit Button Handler Progress Feedback
# ============================================================================

print("\n2. ISSUE 2 - STREAMLIT BUTTON PROGRESS FEEDBACK")
print("-" * 60)

try:
    # Read app.py and verify button handler has progress feedback
    with open("app.py", "r", encoding="utf-8") as f:
        app_content = f.read()

    check("app.py readable", True)

    # Check for required elements in button handler
    button_checks = [
        ("Button has 'Generate Rewritten Resume' text",
         'Generate Rewritten Resume' in app_content),
        ("status_placeholder = st.empty()",
         "status_placeholder = st.empty()" in app_content),
        ("progress_bar = st.progress(0)",
         "progress_bar = st.progress(0)" in app_content),
        ("st.spinner() wrapper",
         "with st.spinner" in app_content and "Generating your upgraded resume" in app_content),
        ("Step 1/4 status message",
         "Step 1/4" in app_content),
        ("Step 2/4 status message",
         "Step 2/4" in app_content),
        ("Step 3/4 status message",
         "Step 3/4" in app_content),
        ("Step 4/4 status message",
         "Step 4/4" in app_content),
        ("Progress bar advances to 100",
         "progress_bar.progress(100)" in app_content),
        ("Success message displayed",
         "status_placeholder.success" in app_content),
        ("Error handling with st.error",
         "st.error(f\"" in app_content and "Resume generation failed" in app_content),
        ("st.exception(e) for traceback",
         "st.exception(e)" in app_content),
        ("st.stop() on error",
         "st.stop()" in app_content),
    ]

    for check_name, result in button_checks:
        check(f"  {check_name}", result)

except Exception as e:
    logger.error(f"Error checking app.py: {e}")
    check("app.py button handler verification", False)

# ============================================================================
# Summary
# ============================================================================

print("\n" + "=" * 60)
print(f"PASSED: {len(passes)} checks")
print(f"FAILED: {len(failures)} checks")

if failures:
    print("\nFailed checks:")
    for failure in failures:
        print(f"  - {failure}")
    sys.exit(1)
else:
    print("\nAll checks passed! Fixes verified.")
    sys.exit(0)
