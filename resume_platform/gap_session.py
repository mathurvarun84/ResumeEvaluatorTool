"""Interactive gap-closing helper.

This module is purely terminal-interface and document-generation logic.
It does not touch any global state other than writing the requested ``.docx``.

Functions
---------
``run_gap_session``
    Orchestrates the interactive diff view, user decisions and the final
    ``.docx`` export.

The behaviour is dictated by the unit-test in ``test_compliance.py`` and by
the architectural rules listed in ``CLAUDE.md``.
"""

from __future__ import annotations

from dataclasses import dataclass
import os
import subprocess
import tempfile
from typing import Dict, List

# Rich and python-docx are required by the project dependencies.
from rich.console import Console
from rich.table import Table
from rich.panel import Panel
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

console = Console()

@dataclass
class SectionResult:
    name: str
    original: str
    rewritten: str
    changes_made: List[str]
    keywords_added: List[str]
    decision: str = ""

# ---------------------------------------------------------------------------
# Helper utilities
# ---------------------------------------------------------------------------

def _add_horizontal_rule(document: Document) -> None:
    """Insert a horizontal line using an XML border.

    The function mutates the last paragraph, adding a bottom border to make a
    visible HR in the generated r-sum-.
    """
    if Document is None:
        return
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # Create a border element.
    p_pr = p._element.get_or_add_pPr()
    border = OxmlElement("w:bottom")
    border.set(qn("w:val"), "single")
    border.set(qn("w:sz"), "3")
    border.set(qn("w:space"), "1")
    border.set(qn("w:color"), "auto")
    border.set(qn("w:themeColor"), "accent1")
    border_element = OxmlElement("w:pBdr")
    border_element.append(border)
    p_pr.append(border_element)

# ---------------------------------------------------------------------------
# Core public function
# ---------------------------------------------------------------------------

def run_gap_session(
    gap_result: Dict,
    original_resume_text: str,
    output_path: str,
    score_delta: Dict | None = None,
) -> Dict:
    """Run the interactive gap-closing session.

    Parameters
    ----------
    gap_result:
        Dictionary in the shape produced by :func:`agents.gap_analyzer`.
    original_resume_text:
        Full raw resume text. Used only to parse the candidate name and contact
        line for the ``.docx`` export.
    output_path:
        Location where the resulting r-sum- should be written.
    score_delta:
        Optional dict with score improvement data:
        {score_before, score_after, delta, keywords_added, sections_improved, remaining_gaps, manual_suggestions}

    Returns
    -------
    dict
        ``{ decisions, output_path, sections_accepted, sections_rejected }``.

    Remarks
    -------
    * All user prompts are issued via :func:`input`. The tests monkeypatch
      :func:`builtins.input` to automate ``'A'`` responses.
    * Editing a section uses the editor from ``$EDITOR`` or ``$VISUAL`` and
      falls back to ``nano`` - ``vi``. If all fail, the section is accepted
      with a warning.
    * The function is written to be safe in both interactive and non-
      interactive environments. It prints only minimal status messages.
    """

    console.print("\n[bold]Gap Closing Session[/bold]")

    console.print()
    before = gap_result.get("jd_match_score_before", 0)
    after = gap_result.get("jd_match_score_after", 0)
    console.print(Panel(f"JD Match: {before}% -> {after}% after all rewrites", style="cyan"))

    sections_results: List[SectionResult] = []
    decisions: Dict[str, str] = {}

    # Handle empty sections gracefully. If no sections were returned from the analyzer, notify user and skip the interactive loop.
    for s in gap_result.get("sections", []):
        sec = SectionResult(
            name=s.get("section_name", "").strip(),
            original=s.get("original", ""),
            rewritten=s.get("rewritten", ""),
            changes_made=s.get("changes_made", []),
            keywords_added=s.get("keywords_added", []),
        )

        console.print(f"\n[bold]{sec.name}[/bold]")

        table = Table(show_header=True, header_style="bold magenta")
        table.add_column("ORIGINAL", style="red")
        table.add_column("REWRITTEN", style="green", overflow="fold")
        table.add_row(sec.original or "(empty)", sec.rewritten or "(empty)")
        console.print(table)

        if sec.changes_made:
            console.print("[yellow]Changes made:[/yellow]")
            for c in sec.changes_made:
                console.print(f" - {c}")

        if sec.keywords_added:
            kw_text = " ".join(f"[yellow]{kw}[/yellow]" for kw in sec.keywords_added)
            console.print(f"[yellow]Keywords added:[/yellow] {kw_text}")

        # Prompt loop
        while True:
            console.print("[bold][A]ccept[/bold] / [bold][R]eject[/bold] / [bold][E]dit[/bold] this section? ", end="")
            choice = input().strip().upper()
            if choice == "A":
                sec.decision = "accepted"
                break
            if choice == "R":
                sec.decision = "rejected"
                break
            if choice == "E":
                edited_text = _edit_section(sec.rewritten)
                sec.rewritten = edited_text
                sec.decision = "edited"
                break
            console.print("[red]Invalid choice, please enter A, R or E.[/red]")

        decisions[sec.name] = sec.decision
        sections_results.append(sec)

    # Summary table
    console.print("\n[bold]Summary[/bold]")
    summary = Table(show_header=True, header_style="bold magenta")
    summary.add_column("Section")
    summary.add_column("Decision")
    summary.add_column("Keywords Added")
    for sec in sections_results:
        kw = ", ".join(sec.keywords_added)
        summary.add_row(sec.name, sec.decision, kw)
    console.print(summary)
    console.print()
    console.print(f"JD Match before: {before}% | after: {after}%")

    # Build .docx with comprehensive metadata
    try:
        _export_to_docx(
            original_resume_text,
            sections_results,
            output_path,
            gap_result=gap_result,
            score_delta=score_delta,
        )
    except Exception as e:
        console.print(f"[red]Error during DOCX export: {e}[/red]")
        import traceback
        traceback.print_exc()
        raise

    # Count
    accepted = sum(1 for d in decisions.values() if d != "rejected")
    rejected = sum(1 for d in decisions.values() if d == "rejected")

    return {
        "decisions": decisions,
        "output_path": output_path,
        "sections_accepted": accepted,
        "sections_rejected": rejected,
    }


# ---------------------------------------------------------------------------
# Internal helpers
# ---------------------------------------------------------------------------

def _edit_section(initial_text: str) -> str:
    """Open an editor for the user to modify the provided text.

    Parameters
    ----------
    initial_text:
        The current rewritten section.

    Returns
    -------
    str
        The edited text, or the original if the user aborts.
    """
    with tempfile.NamedTemporaryFile(mode="w", suffix=".txt", delete=False, encoding="utf-8") as tf:
        tf.write(initial_text)
        tmp_path = tf.name

    # Determine editor candidates: $EDITOR, $VISUAL, or fallbacks.
    editor = os.environ.get("EDITOR") or os.environ.get("VISUAL") or "nano"
    candidates = [editor, "nano", "vi"]  # try fallbacks if earlier missing
    for ed in candidates:
        try:
            result = subprocess.run([ed, tmp_path], check=True)
            if result.returncode == 0:
                # Editor exited successfully - read the file.
                with open(tmp_path, "r", encoding="utf-8") as f:
                    edited = f.read().strip()
                os.unlink(tmp_path)
                return edited
        except FileNotFoundError:
            # Editor binary missing - try next candidate.
            continue
        except subprocess.CalledProcessError as e:
            # Editor failed (non-zero exit) - show warning and keep original.
            console.print(f"[yellow]Editor returned error: {e}[/yellow]")
            os.unlink(tmp_path)
            return initial_text
        except Exception as e:
            # Any other unexpected error - abort editing.
            console.print(f"[red]Editor failed: {e}. Accepting original text.[/red]")
            os.unlink(tmp_path)
            return initial_text

    # All candidates failed.
    console.print("[red]No suitable editor found. Accepting original text.[/red]")
    os.unlink(tmp_path)
    return initial_text


def _export_to_docx(
    full_text: str,
    sections: List[SectionResult],
    path: str,
    gap_result: Dict | None = None,
    score_delta: Dict | None = None,
) -> None:
    """Create a comprehensive r-sum- document with gap analysis metadata.

    The structure includes:
    1. Candidate Name + Target Role (header)
    2. Original Resume Summary (before rewrite) - optional
    3. Rewritten Resume (full - all sections)
    4. Gap Analysis Table (skill/keyword gaps identified)
    5. JD Match Score - Before (with breakdown)
    6. JD Match Score - After (with breakdown)
    7. Delta Commentary (what changed and why)

    Parameters
    ----------
    full_text : str
        Original resume text.
    sections : List[SectionResult]
        Sections with accept/reject/edit decisions.
    path : str
        Output file path.
    gap_result : Dict, optional
        Gap analysis data from orchestrator.
    score_delta : Dict, optional
        Score improvement tracking data.
    """
    import logging
    logger = logging.getLogger(__name__)

    # ISSUE 1 FIX: Diagnostic logging of session data
    logger.debug("=== DOCX GEN SESSION DATA ===")
    logger.debug(f"full_text length: {len(full_text) if full_text else 0}")
    logger.debug(f"sections count: {len(sections) if sections else 0}")
    if gap_result:
        for k, v in gap_result.items():
            v_str = str(v)[:200] if v else "None"
            logger.debug(f"  gap_result[{k}]: {type(v).__name__} = {v_str}")
    if score_delta:
        for k, v in score_delta.items():
            v_str = str(v)[:200] if v else "None"
            logger.debug(f"  score_delta[{k}]: {type(v).__name__} = {v_str}")

    if Document is None:  # pragma: no cover - not needed for tests
        console.print("[red]python-docx not available; cannot export r-sum-.[/red]")
        return
    document = Document()

    # Section 1: Header - name and contact
    # ISSUE 1 FIX: Null guards with fallback text
    lines = [line.strip() for line in full_text.splitlines() if line.strip()] if full_text else []
    name = lines[0] if lines else "Candidate (Name not found)"
    contact = lines[1] if len(lines) > 1 else "Contact information not provided"
    target_role = (gap_result.get("target_role", "") or "Target Role not specified") if gap_result else "Target Role not specified"

    p_name = document.add_heading(name, level=1)
    p_name.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in p_name.runs:
        run.font.size = Pt(16)
        run.font.bold = True

    if target_role:
        p_role = document.add_paragraph(f"Target Role: {target_role}")
        p_role.alignment = WD_ALIGN_PARAGRAPH.CENTER

    if contact:
        p_contact = document.add_paragraph(contact)
        p_contact.alignment = WD_ALIGN_PARAGRAPH.CENTER

    _add_horizontal_rule(document)

    # Section 2: Original Resume Summary (before rewrite) - optional summary
    if gap_result and gap_result.get("show_original_summary"):
        document.add_heading("ORIGINAL RESUME SUMMARY", level=2)
        orig_summary = gap_result.get("original_summary", "")
        if orig_summary:
            document.add_paragraph(orig_summary)
        _add_horizontal_rule(document)

    # Section 3: Rewritten Resume (full - all sections, not truncated)
    # ISSUE 1 FIX: Unconditional section writing - write even if empty
    document.add_heading("UPGRADED RESUME", level=2)
    _add_horizontal_rule(document)

    from engine.resume_builder import build_final_docx as _build_docx
    from parser import _build_structured_resume

    parsed_resume_structured = _build_structured_resume(full_text or "")
    _rewrites = {}
    for sec in sections:
        final_text = sec.original if sec.decision == "rejected" else sec.rewritten
        if final_text:
            _rewrites[sec.name] = {
                "balanced": final_text,
                "aggressive": final_text,
                "top_1_percent": final_text,
            }
    _docx_bytes = _build_docx(
        structured=parsed_resume_structured,
        rewrites=_rewrites,
        style="balanced",
    ) if sections else b""

    if not sections:
        document.add_paragraph("-- No resume sections available for rewrite.")
    else:
        for sec in sections:
            # sec is a SectionResult dataclass with: name, original, rewritten, decision, changes_made, keywords_added
            decision = sec.decision
            if decision == "rejected":
                final_text = sec.original
            else:
                final_text = sec.rewritten

            # ISSUE 1 FIX: Write section with fallback even if empty
            final_text = final_text or f"-- {sec.name} content not available"

            # Section heading - use black ruled border per CLAUDE.md
            p_head = document.add_heading(sec.name.upper(), level=2)
            p_head.alignment = WD_ALIGN_PARAGRAPH.LEFT
            for run in p_head.runs:
                run.font.size = Pt(13)
                run.font.bold = True
                run.font.color.rgb = RGBColor(0, 0, 0)  # black, not teal

            # Add border line
            p_pr = p_head._element.get_or_add_pPr()
            border = OxmlElement("w:bottom")
            border.set(qn("w:val"), "single")
            border.set(qn("w:sz"), "3")
            border.set(qn("w:space"), "1")
            border.set(qn("w:color"), "000000")  # black
            border_element = OxmlElement("w:pBdr")
            border_element.append(border)
            p_pr.append(border_element)

            # Content - split into paragraphs and handle bullets
            # ISSUE 1 FIX: Always write content, even if placeholder text
            content_lines = final_text.splitlines() if isinstance(final_text, str) else []
            if not content_lines:
                document.add_paragraph("(No content)")
            else:
                for line in content_lines:
                    line = line.rstrip()
                    if not line:
                        # Empty line - add spacing paragraph
                        document.add_paragraph()
                        continue
                    # Check if it's a bullet point
                    if line.lstrip().startswith(("-", "-", "*")):
                        # Remove bullet markers and re-add as styled bullet
                        cleaned = line.lstrip("--* ").strip()
                        paragraph = document.add_paragraph(cleaned, style="List Bullet")
                        for run in paragraph.runs:
                            run.font.size = Pt(11)
                    else:
                        # Normal paragraph
                        paragraph = document.add_paragraph(line)
                        for run in paragraph.runs:
                            run.font.size = Pt(11)
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Section 4: Gap Analysis Table (skill/keyword gaps identified)
    # ISSUE 1 FIX: Write gap analysis section unconditionally
    document.add_paragraph()
    _add_horizontal_rule(document)
    document.add_heading("GAP ANALYSIS", level=2)

    # Build gap table
    table = document.add_table(rows=1, cols=4)
    table.style = "Light Grid Accent 1"
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "Section"
    hdr_cells[1].text = "Gap Type"
    hdr_cells[2].text = "Severity"
    hdr_cells[3].text = "Suggested Fix"

    gaps = (gap_result or {}).get("gaps", []) if gap_result else []
    if gaps:
        for gap in gaps:
            row_cells = table.add_row().cells
            row_cells[0].text = str(gap.get("section", "N/A"))
            row_cells[1].text = str(gap.get("type", "N/A"))
            row_cells[2].text = str(gap.get("severity", "N/A"))
            row_cells[3].text = str(gap.get("suggestion", "N/A"))[:100]  # truncate for table
    else:
        # Add placeholder row if no gaps
        row_cells = table.add_row().cells
        row_cells[0].text = "-"
        row_cells[1].text = "No gaps identified"
        row_cells[2].text = "-"
        row_cells[3].text = "Resume is well-aligned with JD"

    # Section 5 & 6: JD Match Scores - Before and After
    # ISSUE 1 FIX: Write JD match analysis unconditionally with fallbacks
    document.add_paragraph()
    _add_horizontal_rule(document)
    document.add_heading("JD MATCH ANALYSIS", level=2)

    # Extract scores safely with fallback values
    before_score = (gap_result or {}).get("jd_match_score_before", 0) if gap_result else 0
    after_score = (gap_result or {}).get("jd_match_score_after", 0) if gap_result else 0

    document.add_heading("BEFORE", level=3)
    document.add_paragraph(f"JD Match Score: {before_score}%")

    document.add_heading("AFTER", level=3)
    document.add_paragraph(f"JD Match Score: {after_score}%")

    delta = after_score - before_score
    if delta != 0:
        document.add_paragraph(f"Improvement: +{delta}%" if delta > 0 else f"Change: {delta}%")
    elif gap_result:
        document.add_paragraph("Score remained the same - focus on the gaps listed above.")

    # Section 7: Delta Commentary (what changed and why the score moved)
    # ISSUE 1 FIX: Write improvement summary unconditionally with fallbacks
    document.add_paragraph()
    _add_horizontal_rule(document)
    document.add_heading("IMPROVEMENT SUMMARY", level=2)

    if score_delta:
        keywords_added = score_delta.get("keywords_added", []) or []
        if keywords_added:
            document.add_heading("Keywords Added:", level=3)
            for kw in keywords_added[:20]:  # limit to top 20
                document.add_paragraph(str(kw), style="List Bullet")
        else:
            document.add_heading("Keywords Added:", level=3)
            document.add_paragraph("No new keywords added.", style="List Bullet")

        sections_improved = score_delta.get("sections_improved", []) or []
        if sections_improved:
            document.add_heading("Sections Strengthened:", level=3)
            for sec in sections_improved:
                document.add_paragraph(str(sec), style="List Bullet")
        else:
            document.add_heading("Sections Strengthened:", level=3)
            document.add_paragraph("No sections were strengthened.", style="List Bullet")

        remaining_gaps = score_delta.get("remaining_gaps", []) or []
        if remaining_gaps:
            document.add_heading("Remaining Gaps:", level=3)
            for gap in remaining_gaps:
                document.add_paragraph(str(gap), style="List Bullet")
        else:
            document.add_heading("Remaining Gaps:", level=3)
            document.add_paragraph("All gaps have been addressed.", style="List Bullet")

        manual_suggestions = score_delta.get("manual_suggestions", []) or []
        if manual_suggestions:
            document.add_heading("Manual Suggestions:", level=3)
            for sugg in manual_suggestions:
                document.add_paragraph(str(sugg), style="List Bullet")
        else:
            document.add_heading("Manual Suggestions:", level=3)
            document.add_paragraph("No additional suggestions.", style="List Bullet")
    else:
        document.add_paragraph("-- No improvement data available.", style="List Bullet")

    # Save - ensure parent directory exists
    parent = os.path.dirname(os.path.abspath(path))
    if parent and not os.path.exists(parent):
        os.makedirs(parent, exist_ok=True)
    document.save(path)

    # ISSUE 1 FIX: Assertion to verify DOCX was not blank
    output_size = os.path.getsize(path) if os.path.exists(path) else 0
    logger.debug(f"DOCX saved - size: {output_size} bytes")
    assert output_size > 5000, (
        f"DOCX generation likely failed - output is only {output_size} bytes "
        "(blank DOCX is ~4KB). Check session data above."
    )

    console.print(f"[green]Upgraded r-sum- saved to {path}[/green]")

